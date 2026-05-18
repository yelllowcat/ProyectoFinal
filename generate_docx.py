#!/usr/bin/env python3
"""
Generates the IEEE-830 style documentation for the UNIRED project as .docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime

OUTPUT = "/media/david/Storage/repos/ProyectoFinal/DOCUMENTACION_UNIRED.docx"

DARK_BLUE = RGBColor(0x1a, 0x3a, 0x5c)
MED_BLUE = RGBColor(0x2c, 0x5f, 0x8a)
HEADER_BG = "1a3a5c"
TABLE_HEADER_BG = "2c5f8a"
ALT_ROW_BG = "f4f7fb"


# ── Helper functions ───────────────────────────────────────────
def set_cell_shading(cell, color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    tcPr.append(shading)


def add_hr(doc):
    """Add a horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{HEADER_BG}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def heading(doc, text, level=1):
    """Add a section heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE if level == 1 else MED_BLUE
        run.font.name = 'Times New Roman'
    return h


def body(doc, text, bold=False, italic=False, alignment=None):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    return p


def bullet(doc, text):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    return p


def make_table(doc, data, col_widths=None, header=True):
    """Create a styled table."""
    rows_count = len(data)
    cols_count = len(data[0]) if data else 0
    table = doc.add_table(rows=rows_count, cols=cols_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            if i == 0 and header:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(cell, TABLE_HEADER_BG)
            elif i % 2 == 0 and not header:
                set_cell_shading(cell, ALT_ROW_BG)
            elif i % 2 != 0 and header:
                pass

            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()  # spacer after table
    return table


def make_kv_table(doc, data):
    """Key-Value table with first column as header."""
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (key, val) in enumerate(data):
        # Key cell
        cell_key = table.rows[i].cells[0]
        cell_key.text = ''
        p = cell_key.paragraphs[0]
        run = p.add_run(str(key))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell_key, TABLE_HEADER_BG)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        # Value cell
        cell_val = table.rows[i].cells[1]
        cell_val.text = ''
        p = cell_val.paragraphs[0]
        run = p.add_run(str(val))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    for row in table.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.3)

    doc.add_paragraph()
    return table


def page_break(doc):
    doc.add_page_break()


# ── Content Builders ──────────────────────────────────────────

def build_cover(doc):
    """Build the cover page."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD AUTÓNOMA DE BAJA CALIFORNIA SUR")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Departamento Académico de Sistemas Computacionales")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("La Paz, Baja California Sur")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

    doc.add_paragraph()
    add_hr(doc)
    doc.add_paragraph()

    labels = [
        ("ASIGNATURA", "Ingeniería de Software II"),
        ("SEMESTRE", "6° Semestre (T.V.)"),
        ("DOCENTE", "Italia Estrada Cota"),
    ]
    for label, value in labels:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INTEGRANTES")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = True
    for name in ["Gabriel Lauro Hernández", "Jose Manuel Orozco Vazquez", "David Gonzalez Vargas"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    doc.add_paragraph()
    add_hr(doc)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SISTEMA INTEGRAL DE RED SOCIAL UNIVERSITARIA")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIRED")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.font.color.rgb = MED_BLUE

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DOCUMENTACIÓN TÉCNICA Y DE ANÁLISIS")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

    for info in ["Versión: 1.0", "Mayo 2026"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(info)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    page_break(doc)


def build_feasibility(doc):
    """Section 1: Estudio de Factibilidad"""
    heading(doc, "1. ESTUDIO DE FACTIBILIDAD", 1)
    body(doc, (
        "Durante la fase inicial del proyecto, se realizó una investigación exhaustiva dentro de la comunidad "
        "universitaria para confirmar que el sistema solicitado (UNIRED) puede desarrollarse e implementarse sin "
        "problemas. Esta investigación ayudó a verificar que el nuevo sistema realmente ayudará a las necesidades "
        "actuales de los estudiantes, facilitando los procesos de comunicación, interacción académica y social. "
        "Además, se revisó que el sistema puede integrarse con los equipos existentes, que el personal está dispuesto "
        "a utilizarlo y que los costos de desarrollo son manejables. El análisis se dividió en tres áreas principales: "
        "factibilidad técnica, operativa y económica."
    ))

    # 1.1 Technical
    heading(doc, "1.1 Factibilidad técnica", 2)
    body(doc, (
        "La evaluación técnica se realizó con base en los recursos tecnológicos que actualmente utiliza la comunidad "
        "y los que se planean incorporar durante la implementación del sistema."
    ))

    heading(doc, "1.1.1 Hardware", 3)
    body(doc, "Se verificó que los dispositivos móviles y computadoras utilizadas por los estudiantes cumplen con los requisitos mínimos para ejecutar el sistema web de UNIRED:")
    bullet(doc, "Equipos con procesador de doble núcleo o superior.")
    bullet(doc, "Memoria RAM de al menos 4 GB para una navegación fluida.")
    bullet(doc, "Conexión a internet estable (3G/4G o Wi-Fi).")
    bullet(doc, "Servidor de base de datos con al menos 20 GB de almacenamiento SSD, escalable a 100 GB.")
    body(doc, "Para alojar el sistema se optó por un servidor en la nube (Hostinger), el cual permite acceso seguro y disponibilidad constante.")

    heading(doc, "1.1.2 Software", 3)
    body(doc, "Durante la reunión técnica del equipo de desarrollo, se acordó implementar la siguiente arquitectura y stack tecnológico:")
    bullet(doc, "Back-end: PHP 7.4+, por su facilidad de mantenimiento, escalabilidad en entornos compartidos y gran soporte comunitario.")
    bullet(doc, "Front-end: HTML5, CSS3, y JavaScript (Vanilla) enfocado en un diseño minimalista, moderno y responsive.")
    bullet(doc, "Base de datos: MySQL, debido a su estructura relacional sólida que se adapta perfectamente al manejo de usuarios, publicaciones, comentarios y relaciones de amistad.")
    bullet(doc, "Servidor Web: Apache con mod_rewrite habilitado.")
    bullet(doc, "Dependencias: Composer, vlucas/phpdotenv para gestión de variables de entorno, y TCPDF para generación de reportes en PDF.")
    bullet(doc, "Gráficos: Chart.js v4 para visualización de estadísticas en el dashboard administrativo.")
    body(doc, "El equipo definió que, al finalizar el proyecto, se entregará la documentación técnica, manual de usuario y el sistema desplegado en producción.")

    heading(doc, "1.1.3 Personal técnico requerido", 3)
    make_table(doc, [
        ["Rol", "Cantidad"],
        ["Analista de Sistemas", "1"],
        ["Diseñador UI/UX", "1"],
        ["Desarrollador Backend PHP", "2"],
        ["Desarrollador Frontend", "2"],
        ["Administrador de Base de Datos (DBA)", "1"],
        ["Ingeniero de Pruebas (QA)", "1"],
    ], [3.0, 1.5])

    # 1.2 Operational
    heading(doc, "1.2 Factibilidad operativa", 2)
    body(doc, (
        "El sistema UNIRED será adoptado fácilmente por la comunidad estudiantil, ya que actualmente experimentan "
        "dificultades para centralizar la comunicación universitaria en otras plataformas genéricas. La carga de gestión "
        "de contactos disminuirá gracias a la automatización y búsqueda de perfiles dentro del campus. Entre los "
        "beneficios operativos identificados se encuentran:"
    ))
    bullet(doc, "La comunicación universitaria estará centralizada en un solo lugar exclusivo para la comunidad.")
    bullet(doc, "Los usuarios muestran una altísima disposición a usar el sistema debido a la interfaz intuitiva y los esquemas de colores amigables.")
    bullet(doc, "Se utilizará la tipografía Roboto de Google Fonts con pesos 300, 400 y 700 para garantizar jerarquía visual.")
    bullet(doc, "El color primario Azul Turquesa (#4db8c4) combinado con blancos limpios y grises oscuros brinda una experiencia agradable.")
    bullet(doc, "La interfaz responsive se adapta a dispositivos móviles (320px hasta 1920px).")
    bullet(doc, "Las microinteracciones incluyen transiciones de 0.3s ease-in-out en botones y esqueletos de carga durante peticiones asíncronas.")
    bullet(doc, "Se ofrece un dashboard administrativo con métricas y gráficas para la toma de decisiones.")

    # 1.3 Economic
    heading(doc, "1.3 Factibilidad económica", 2)
    body(doc, "Se realizó una estimación del costo total del proyecto comparado con los beneficios que generará para la comunidad universitaria.")
    make_table(doc, [
        ["Concepto", "Costo aproximado"],
        ["Desarrollo del software (equipo de 7 personas)", "$175,000"],
        ["Servidor en la nube, dominio y base de datos", "$6,000"],
        ["Soporte técnico anual", "$12,000"],
        ["Capacitación", "$3,000"],
        ["Total", "$196,000"],
    ], [3.5, 2.0])

    body(doc, "Beneficios esperados:", bold=True)
    body(doc, "Durante las entrevistas con la comunidad universitaria se determinaron beneficios reales como:")
    bullet(doc, "Comunicación centralizada y exclusiva para la comunidad universitaria.")
    bullet(doc, "Reducción de la dispersión de información en múltiples plataformas (WhatsApp, Facebook, correo).")
    bullet(doc, "Sistema de publicaciones con soporte multimedia y gestión de contenido.")
    bullet(doc, "Red de amistades y contactos gestionable desde un solo lugar.")
    bullet(doc, "Dashboard administrativo con estadísticas en tiempo real.")
    bullet(doc, "Seguridad en los datos y control de acceso basado en roles.")
    body(doc, (
        "La comunidad universitaria consideró que los beneficios justifican plenamente la inversión, pues la plataforma "
        "permitirá mayor conexión entre estudiantes, docentes y personal administrativo, mejorando la experiencia "
        "universitaria integral."
    ))

    page_break(doc)


def build_interviews(doc):
    """Section 2: Recopilación de Información - Entrevistas"""
    heading(doc, "2. RECOPILACIÓN DE INFORMACIÓN: ENTREVISTA", 1)

    heading(doc, "2.1 Personas entrevistadas", 2)
    make_table(doc, [
        ["Puesto", "Motivo de selección"],
        ["Estudiante Universitario", "Conoce de primera mano las necesidades de comunicación, interacción social y académica dentro del campus."],
        ["Docente Universitario", "Aporta información sobre la necesidad de espacios estructurados para compartir conocimiento y materiales académicos."],
        ["Administrador de Sistemas", "Responsable de la seguridad, almacenamiento de datos y escalabilidad de la plataforma."],
    ], [1.8, 4.7])
    body(doc, "Estas tres perspectivas permitieron obtener un panorama general del funcionamiento y las necesidades de comunicación dentro de la comunidad universitaria.")

    heading(doc, "2.2 Guión de preguntas realizadas", 2)

    heading(doc, "2.2.1 Preguntas para el Estudiante Universitario", 3)
    q_student = [
        "1. ¿Cuáles son las mayores dificultades que encuentra actualmente en la comunicación universitaria?",
        "2. ¿Qué funcionalidades consideraría indispensables en una red social universitaria?",
        "3. ¿Qué nivel de privacidad espera del sistema?",
        "4. ¿Cómo preferiría gestionar sus contactos y amistades dentro de la plataforma?",
        "5. ¿Qué tipo de contenido le gustaría compartir en sus publicaciones?",
        "6. ¿Considera útil un sistema de comentarios y reacciones en las publicaciones?",
        "7. ¿Qué tan importante es para usted la búsqueda de otros estudiantes por nombre o carrera?",
        "8. ¿Le gustaría poder personalizar su perfil con foto, biografía y datos académicos?",
    ]
    for q in q_student:
        bullet(doc, q)

    heading(doc, "2.2.2 Preguntas para el Docente Universitario", 3)
    q_teacher = [
        "1. ¿Cuáles son las mayores dificultades que encuentra actualmente en la comunicación universitaria?",
        "2. ¿Qué funcionalidades consideraría indispensables en una red social universitaria?",
        "3. ¿Considera necesario un espacio donde los alumnos compartan conocimiento de manera estructurada?",
        "4. ¿Qué tipo de supervisión o moderación le gustaría tener sobre el contenido publicado?",
        "5. ¿Cree útil poder crear grupos o foros por materia o carrera?",
        "6. ¿Qué información académica debería mostrarse en los perfiles de los estudiantes?",
        "7. ¿Considera importante la integración con otras herramientas académicas?",
        "8. ¿Qué nivel de seguridad y privacidad espera para los datos de los estudiantes?",
    ]
    for q in q_teacher:
        bullet(doc, q)

    heading(doc, "2.2.3 Preguntas para el Administrador de Sistemas", 3)
    q_admin = [
        "1. ¿Cuáles son las mayores dificultades que encuentra actualmente en la comunicación universitaria?",
        "2. ¿Qué funcionalidades consideraría indispensables en una red social universitaria?",
        "3. ¿Qué nivel de privacidad y seguridad espera del sistema?",
        "4. ¿Cómo debe gestionarse el almacenamiento de imágenes de perfil y publicaciones?",
        "5. ¿Qué consideraciones de escalabilidad deben tomarse para la base de datos MySQL?",
        "6. ¿Qué métricas y estadísticas considera necesarias para monitorear la plataforma?",
        "7. ¿Debe el sistema permitir la exportación de reportes en PDF?",
        "8. ¿Qué tan importante es la seguridad y respaldo de la información?",
    ]
    for q in q_admin:
        bullet(doc, q)

    heading(doc, "2.3 Resultados obtenidos de la entrevista", 2)

    heading(doc, "2.3.1 Estudiante Universitario", 3)
    bullet(doc, "Actualmente se utilizan herramientas muy dispersas para la comunicación universitaria (WhatsApp, Facebook, correo electrónico).")
    bullet(doc, "No existe un punto centralizado que sea exclusivo para la universidad.")
    bullet(doc, "Desea un lugar seguro donde todos los miembros sean parte verificada de la comunidad.")
    bullet(doc, "Considera indispensable un sistema de publicaciones con capacidad de subir fotos, dar \"me gusta\" y comentar.")
    bullet(doc, "Requiere un buscador de usuarios eficiente por nombre o carrera.")
    bullet(doc, "Es crucial la gestión de solicitudes de amistad (enviar, aceptar, rechazar).")
    bullet(doc, "Necesita poder editar su perfil de forma segura (nombre, biografía, foto).")
    bullet(doc, "Desea poder eliminar su cuenta si así lo decide, con modales de confirmación para evitar accidentes.")

    heading(doc, "2.3.2 Docente Universitario", 3)
    bullet(doc, "Actualmente la comunicación con alumnos se dispersa en múltiples canales no oficiales.")
    bullet(doc, "Desea un espacio ordenado donde los alumnos puedan compartir conocimiento de manera estructurada.")
    bullet(doc, "Considera indispensable un sistema de publicaciones donde se pueda compartir contenido académico.")
    bullet(doc, "Valora la posibilidad de crear grupos o foros por materia para discusiones académicas.")
    bullet(doc, "Requiere que el sistema tenga control de roles y permisos para moderación de contenido.")
    bullet(doc, "Espera un alto nivel de seguridad y privacidad para los datos de los estudiantes.")

    heading(doc, "2.3.3 Administrador de Sistemas", 3)
    bullet(doc, "Existe una preocupación importante por la seguridad en el manejo de datos personales.")
    bullet(doc, "Se requiere un control de almacenamiento eficiente para imágenes de perfil y publicaciones (límites de tamaño).")
    bullet(doc, "La base de datos MySQL debe estar optimizada para consultas rápidas mediante Stored Procedures.")
    bullet(doc, "Se necesita un dashboard con métricas como: usuarios registrados, publicaciones, comentarios, actividad diaria.")
    bullet(doc, "El sistema debe permitir la generación de reportes en PDF con estadísticas relevantes.")
    bullet(doc, "La arquitectura debe seguir el patrón MVC para facilitar el mantenimiento y la escalabilidad.")
    bullet(doc, "Debe implementarse autenticación basada en sesiones con contraseñas hasheadas (bcrypt).")
    bullet(doc, "Todas las entradas de usuario deben ser sanitizadas para prevenir inyección SQL y XSS.")

    heading(doc, "2.4 Conclusiones de la técnica", 2)
    body(doc, "La entrevista permitió identificar problemas de comunicación y necesidades importantes:")
    bullet(doc, "Los procesos de comunicación actuales están completamente dispersos en múltiples plataformas genéricas.")
    bullet(doc, "No existe un espacio digital centralizado y exclusivo para la comunidad universitaria.")
    bullet(doc, "Los estudiantes, docentes y personal administrativo requieren un sistema que organice publicaciones, comentarios, gestión de amistades y perfiles.")
    bullet(doc, "Se identificó la necesidad de roles diferenciados (usuario regular y administrador).")
    bullet(doc, "La comunidad demanda un dashboard administrativo con métricas y estadísticas para la toma de decisiones.")
    bullet(doc, "La seguridad y privacidad de los datos es una prioridad fundamental para todos los actores.")
    body(doc, "La información recopilada confirma la necesidad inmediata de implementar un sistema de red social universitaria que modernice y centralice la comunicación dentro del campus.")

    page_break(doc)


def build_requirements_table(doc):
    """Section 3: Tabla de especificaciones de requerimientos"""
    heading(doc, "3. TABLA DE ESPECIFICACIONES DE REQUERIMIENTOS", 1)
    body(doc, "A continuación se clasifican las preguntas de la entrevista según el tipo de requerimiento (funcional o no funcional):")

    reqs = [
        ["Pregunta", "Tipo de requerimiento"],
        ["¿Cuáles son las mayores dificultades que encuentra actualmente en la comunicación universitaria?", "Funcional"],
        ["¿Qué funcionalidades consideraría indispensables en una red social universitaria?", "Funcional"],
        ["¿Qué nivel de privacidad espera del sistema?", "No funcional"],
        ["¿Cómo preferiría gestionar sus contactos y amistades dentro de la plataforma?", "Funcional"],
        ["¿Qué tipo de contenido le gustaría compartir en sus publicaciones?", "Funcional"],
        ["¿Considera útil un sistema de comentarios y reacciones en las publicaciones?", "Funcional"],
        ["¿Qué tan importante es para usted la búsqueda de otros estudiantes por nombre o carrera?", "Funcional"],
        ["¿Le gustaría poder personalizar su perfil con foto, biografía y datos académicos?", "Funcional"],
        ["¿Considera necesario un espacio donde los alumnos compartan conocimiento de manera estructurada?", "Funcional"],
        ["¿Qué tipo de supervisión o moderación le gustaría tener sobre el contenido publicado?", "Funcional"],
        ["¿Cree útil poder crear grupos o foros por materia o carrera?", "Funcional"],
        ["¿Qué información académica debería mostrarse en los perfiles de los estudiantes?", "Funcional"],
        ["¿Considera importante la integración con otras herramientas académicas?", "No funcional"],
        ["¿Qué nivel de seguridad y privacidad espera para los datos de los estudiantes?", "No funcional"],
        ["¿Cómo debe gestionarse el almacenamiento de imágenes de perfil y publicaciones?", "No funcional"],
        ["¿Qué consideraciones de escalabilidad deben tomarse para la base de datos MySQL?", "No funcional"],
        ["¿Qué métricas y estadísticas considera necesarias para monitorear la plataforma?", "Funcional"],
        ["¿Debe el sistema permitir la exportación de reportes en PDF?", "Funcional"],
        ["¿Qué tan importante es la seguridad y respaldo de la información?", "No funcional"],
        ["¿Cómo debe ser el proceso de registro e inicio de sesión?", "Funcional"],
        ["¿Qué información es obligatoria para el registro de un nuevo usuario?", "Funcional"],
        ["¿Debe el sistema tener soporte para dispositivos móviles?", "No funcional"],
        ["¿Qué velocidad de respuesta espera del sistema?", "No funcional"],
        ["¿Es necesario que el sistema tenga un dashboard administrativo?", "Funcional"],
    ]
    make_table(doc, reqs, [3.5, 2.0])
    page_break(doc)


def build_functional_requirements(doc):
    """Section 4: Tabla de especificación de requerimientos funcionales"""
    heading(doc, "4. TABLA DE ESPECIFICACIÓN DE REQUERIMIENTOS FUNCIONALES", 1)
    body(doc, "En esta sección se detallan los requerimientos funcionales del sistema, especificando qué hace cada uno y qué información necesita.")

    func_reqs = [
        ["¿Qué hace? (Requerimiento funcional)", "¿Qué necesita?"],
        ["Registrar nuevo usuario", "Nombre completo, correo electrónico, contraseña (bcrypt)."],
        ["Iniciar sesión de usuario", "Correo electrónico, contraseña. Validación y regeneración de sesión."],
        ["Cerrar sesión", "Sesión activa. Destrucción de variables de sesión."],
        ["Visualizar perfil de usuario", "ID de usuario. Datos: nombre, biografía, foto, fecha de registro."],
        ["Editar perfil de usuario", "Nombre completo, biografía, foto de perfil (JPEG/PNG/GIF, máximo 5MB)."],
        ["Eliminar cuenta de usuario", "Confirmación de eliminación en cascada de todos los datos."],
        ["Crear publicación", "Contenido de texto, imagen opcional (JPEG/PNG/GIF, máximo 5MB)."],
        ["Editar publicación", "ID de publicación, nuevo contenido de texto y/o imagen."],
        ["Eliminar publicación", "ID de publicación. Verificación de ownership. Soft delete."],
        ["Dar 'Me gusta' a publicación", "ID de publicación, ID de usuario. Toggle like/unlike."],
        ["Agregar comentario a publicación", "ID de publicación, ID de usuario, contenido del comentario."],
        ["Eliminar comentario", "ID de comentario. Verificación de ownership. Soft delete."],
        ["Agregar respuesta a comentario", "ID de comentario, ID de usuario, contenido de la respuesta."],
        ["Eliminar respuesta", "ID de respuesta. Verificación de ownership. Soft delete."],
        ["Dar 'Me gusta' a comentario", "ID de comentario, ID de usuario. Toggle like/unlike."],
        ["Dar 'Me gusta' a respuesta", "ID de respuesta, ID de usuario. Toggle like/unlike."],
        ["Carga paginada de comentarios", "ID de publicación. Carga progresiva de 3 en 3 comentarios."],
        ["Enviar solicitud de amistad", "ID del usuario receptor. Verificación de estado previo."],
        ["Aceptar solicitud de amistad", "ID de solicitud. Cambio de estado a 'accepted'."],
        ["Rechazar solicitud de amistad", "ID de solicitud. Cambio de estado a 'rejected'."],
        ["Cancelar solicitud de amistad", "ID del usuario receptor. Cambio de estado a 'cancelled'."],
        ["Eliminar amigo", "ID de amistad o ID de usuario. Eliminación del registro en friends."],
        ["Buscar usuarios", "Nombre o término de búsqueda. Filtrado por coincidencia."],
        ["Ver lista de amigos", "ID de usuario. Lista de relaciones de amistad confirmadas."],
        ["Ver solicitudes pendientes", "ID de usuario. Solicitudes entrantes con estado 'pending'."],
        ["Sugerencias de usuarios", "ID de usuario. Algoritmo que excluye amigos y solicitudes existentes."],
        ["Dashboard administrativo", "Rol de administrador. Estadísticas y gráficas."],
        ["Gráficas de actividad", "Datos históricos. Chart.js para visualización de timeline y métricas."],
        ["Generar reporte PDF", "Datos estadísticos. TCPDF para generación de reporte descargable."],
        ["Registro de auditoría de perfil", "Datos anteriores y nuevos del perfil. Trigger BEFORE UPDATE."],
        ["Visualización de imágenes en modal", "Ruta de imagen. Modal de pantalla completa con overlay."],
    ]
    make_table(doc, func_reqs, [3.0, 3.5])

    heading(doc, "4.1 Especificación detallada de requerimientos principales (IEEE-830)", 2)

    specs = [
        ("RF-01", "Registro de Usuario", "Funcional",
         "El sistema permitirá a un nuevo estudiante registrarse en la red social universitaria proporcionando sus datos básicos.",
         "Nombre completo, correo electrónico, contraseña.",
         "Validación de formato de correo, hash de contraseña con bcrypt, inserción en BD mediante Stored Procedure sp_register_user.",
         "Confirmación visual de registro exitoso y redirección al inicio de sesión.",
         "Usuario Regular / Estudiante", "No estar autenticado previamente.", "Usuario registrado y redirigido al login."),
        ("RF-02", "Inicio de Sesión", "Funcional",
         "El sistema autenticará a un usuario con sus credenciales (correo y contraseña).",
         "Correo electrónico, contraseña.",
         "Verificación de credenciales mediante sp_login_user, validación con password_verify(), regeneración de ID de sesión.",
         "Redirección al feed de publicaciones (rol user) o dashboard (rol admin).",
         "Usuario Regular / Administrador", "Tener una cuenta registrada y activa.", "Sesión iniciada con variables de sesión."),
        ("RF-03", "Crear Publicación", "Funcional",
         "El sistema permitirá al usuario crear una publicación con texto y opcionalmente una imagen.",
         "Contenido de texto, imagen (opcional, JPEG/PNG/GIF, máximo 5MB).",
         "Validación de tipo y tamaño de archivo, sanitización de texto, inserción mediante sp_create_post.",
         "Publicación renderizada en el feed general con datos del autor y timestamp.",
         "Usuario Regular", "Estar autenticado.", "Post guardado en BD y visible en el feed."),
        ("RF-04", "Sistema de Comentarios", "Funcional",
         "El sistema permitirá agregar comentarios a las publicaciones y visualizarlos con paginación progresiva.",
         "ID de publicación, contenido del comentario.",
         "Inserción mediante sp_create_comment, carga paginada de 3 en 3 mediante AJAX/Fetch.",
         "Comentario visible bajo la publicación con nombre y foto del autor.",
         "Usuario Regular", "Publicación existente y sesión iniciada.", "Comentario agregado al post con paginación actualizada."),
        ("RF-05", "Sistema de Amigos", "Funcional",
         "El sistema gestionará el flujo bidireccional de solicitudes de amistad: enviar, aceptar, rechazar y cancelar.",
         "ID de usuario destino, ID de solicitud.",
         "Validación de estado previo, inserción en friend_requests, actualización de estado, creación en friends al aceptar.",
         "Notificación visual del estado de la solicitud y actualización de listas.",
         "Usuario Regular", "Estar autenticado. No tener solicitud pendiente previa.", "Estado de amistad actualizado en BD."),
        ("RF-06", "Editar Perfil", "Funcional",
         "El sistema permitirá al usuario modificar su nombre completo, biografía y foto de perfil.",
         "Nombre completo, biografía, archivo de imagen (JPEG/PNG/GIF, máximo 5MB).",
         "Validación de campos, carga de imagen a assets/imagesProfile/, actualización en BD, trigger de auditoría.",
         "Perfil actualizado con confirmación visual y registro en user_update_log.",
         "Usuario Regular", "Estar autenticado.", "Datos actualizados en tabla users y renderizados instantáneamente."),
        ("RF-07", "Dashboard Administrativo", "Funcional",
         "El sistema proveerá un panel de control con estadísticas, gráficas y tablas de datos para el administrador.",
         "Rol de administrador, datos agregados de BD.",
         "Consultas a BD mediante AdminModel: getSummaryStats, getActivityTimeline, getEngagementBreakdown, etc.",
         "Dashboard con tarjetas de resumen, gráficas Chart.js y tablas de datos.",
         "Administrador", "Estar autenticado con rol 'admin'.", "Dashboard renderizado con métricas actualizadas."),
        ("RF-08", "Generación de Reporte PDF", "Funcional",
         "El sistema generará un reporte descargable en PDF con las estadísticas más relevantes de la plataforma.",
         "Datos estadísticos de la BD.",
         "Recopilación de datos mediante AdminModel, generación de PDF con TCPDF, headers para descarga.",
         "Archivo PDF descargable con rankings de usuarios y publicaciones.",
         "Administrador", "Estar autenticado con rol 'admin'.", "PDF generado y descargado por el navegador."),
    ]

    for spec in specs:
        data = [
            ("ID del requerimiento", spec[0]),
            ("Nombre o título", spec[1]),
            ("Tipo", spec[2]),
            ("Descripción", spec[3]),
            ("Entradas", spec[4]),
            ("Procesamiento", spec[5]),
            ("Salidas", spec[6]),
            ("Actores involucrados", spec[7]),
            ("Precondiciones", spec[8]),
            ("Postcondiciones", spec[9]),
        ]
        make_kv_table(doc, data)

    page_break(doc)


def build_validation(doc):
    """Section 5: Validación de requerimientos"""
    heading(doc, "5. VALIDACIÓN DE REQUERIMIENTOS", 1)
    body(doc, "En esta sección se validan los requerimientos funcionales del sistema, verificando que sean no ambiguos, relevantes, consistentes y completos.")

    validations = [
        ["Requerimiento", "No ambiguo", "Relevante", "Consistente", "Completo"],
        ["RF-01 — Registro de Usuario", "✓", "✓", "✓", "✓"],
        ["RF-02 — Inicio de Sesión", "✓", "✓", "✓", "✓"],
        ["RF-03 — Crear Publicación", "✓", "✓", "✓", "✓"],
        ["RF-04 — Sistema de Comentarios", "✓", "✓", "✓", "✓"],
        ["RF-05 — Sistema de Amigos", "✓", "✓", "✓", "✓"],
        ["RF-06 — Editar Perfil", "✓", "✓", "✓", "✓"],
        ["RF-07 — Dashboard Administrativo", "✓", "✓", "✓", "✓"],
        ["RF-08 — Generación de Reporte PDF", "✓", "✓", "✓", "✓"],
        ["RF-09 — Eliminar Publicación", "✓", "✓", "✓", "✓"],
        ["RF-10 — Dar Me Gusta a Publicación", "✓", "✓", "✓", "✓"],
        ["RF-11 — Enviar Solicitud de Amistad", "✓", "✓", "✓", "✓"],
        ["RF-12 — Aceptar/Rechazar Amistad", "✓", "✓", "✓", "✓"],
        ["RF-13 — Búsqueda de Usuarios", "✓", "✓", "✓", "✓"],
        ["RF-14 — Carga Paginada de Comentarios", "✓", "✓", "✓", "✓"],
        ["RF-15 — Eliminar Cuenta", "✓", "✓", "✓", "✓"],
        ["RF-16 — Cerrar Sesión", "✓", "✓", "✓", "✓"],
        ["RF-17 — Visualizar Perfil", "✓", "✓", "✓", "✓"],
        ["RF-18 — Sugerencias de Usuarios", "✓", "✓", "✓", "✓"],
        ["RF-19 — Likes en Comentarios", "✓", "✓", "✓", "✓"],
        ["RF-20 — Likes en Respuestas", "✓", "✓", "✓", "✓"],
        ["RF-21 — Auditoría de Cambios de Perfil", "✓", "✓", "✓", "✓"],
        ["RF-22 — Visualización de Imágenes en Modal", "✓", "✓", "✓", "✓"],
        ["RF-23 — Gráficas de Actividad", "✓", "✓", "✓", "✓"],
        ["RF-24 — Diseño Responsive", "✓", "✓", "✓", "✓"],
    ]
    make_table(doc, validations, [2.8, 0.8, 0.8, 0.9, 0.8])
    page_break(doc)


def build_analysis_unit(doc):
    """Section 6: Unidad de Análisis"""
    heading(doc, "6. UNIDAD DE ANÁLISIS", 1)
    body(doc, (
        "La arquitectura de UNIRED se compone de múltiples módulos interconectados, diseñados bajo el patrón MVC "
        "(Modelo-Vista-Controlador) con soporte para operaciones asíncronas mediante AJAX/Fetch y una API RESTful."
    ))

    modules = [
        ("6.1 Módulo de Usuarios",
         "Encargado de gestionar el registro, inicio de sesión seguro y personalización de perfiles. Valida contraseñas "
         "mediante bcrypt, administra sesiones con regeneración de ID, y permite modificar fotos de perfil (hasta 5 MB), "
         "nombre completo y biografía. Incluye un trigger de auditoría (trg_user_update_log) que registra todos los cambios "
         "realizados al perfil en la tabla user_update_log. El módulo también gestiona la eliminación de cuenta con borrado "
         "en cascada de todos los datos asociados (publicaciones, comentarios, likes, amistades)."),
        ("6.2 Módulo de Publicaciones",
         "El corazón de la red social. Permite crear publicaciones de texto y multimedia. Incluye capacidades avanzadas "
         "para subir imágenes con validaciones de tipo (JPEG, PNG, GIF), edición posterior y eliminación segura mediante "
         "soft delete (bandera active). Integra un modal de pantalla completa para visualización de imágenes. Implementa "
         "el sistema de 'Me gusta' con toggle like/unlike y contador en tiempo real. Utiliza la vista v_posts_stats que "
         "combina posts con información del autor, conteo de likes y conteo de comentarios para renderizado eficiente del feed."),
        ("6.3 Módulo de Comentarios",
         "Sub-módulo interactivo que permite añadir respuestas a publicaciones. Utiliza paginación progresiva (de 3 en 3) "
         "para mejorar el rendimiento de la carga inicial, con botón 'Cargar más comentarios'. Incluye enlaces dinámicos a "
         "los perfiles de los autores. Soporta eliminación de comentarios propios (soft delete) y carga dinámica. Integra "
         "un sistema de respuestas anidadas (replies) con funcionalidad similar, y sistema de likes independiente para "
         "comentarios y respuestas."),
        ("6.4 Módulo de Amigos",
         "Gestor de redes de contactos. Soporta flujo bidireccional de solicitudes de amistad (Enviar, Aceptar, Rechazar, "
         "Cancelar). Provee un motor de búsqueda de perfiles públicos por nombre y un algoritmo de sugerencias que excluye "
         "amigos existentes y solicitudes pendientes. Administra la lista de contactos con opción de eliminar amigo. Utiliza "
         "dos tablas: friend_requests para el flujo de solicitudes y friends para las relaciones confirmadas, con constraint "
         "UNIQUE para evitar duplicados."),
        ("6.5 Módulo de Administración",
         "Panel de control exclusivo para usuarios con rol 'admin'. Proporciona un dashboard con 5 tarjetas de resumen "
         "(total usuarios, publicaciones, comentarios, likes, amistades), gráficas interactivas generadas con Chart.js "
         "(activity timeline, engagement breakdown, user growth, user activity split, posts by day of week, post image "
         "ratio, top engaged users) y 4 tablas de datos con los rankings de usuarios y publicaciones más activos. Incluye "
         "generación de reporte PDF descargable mediante TCPDF con las estadísticas más relevantes."),
    ]

    for title, desc in modules:
        heading(doc, title, 2)
        body(doc, desc)

    page_break(doc)


def build_use_cases(doc):
    """Section 7: Documentación - Casos de Uso y Actores"""
    heading(doc, "7. DOCUMENTACIÓN: CASOS DE USO Y ACTORES", 1)
    heading(doc, "7.1 Casos de Uso", 2)
    body(doc, "A continuación se presentan los casos de uso principales del sistema UNIRED, documentados según el formato estándar.")

    # CU-01
    heading(doc, "7.1.1 Caso de Uso — Registro de Usuario", 3)
    make_kv_table(doc, [
        ("Identificador", "CU-01"),
        ("Nombre", "Registro de Usuario"),
        ("Actor Principal", "Estudiante / Usuario Regular"),
        ("Descripción", "Permite a un nuevo estudiante registrarse en la red social universitaria."),
        ("Precondición", "No estar autenticado previamente en el sistema."),
        ("Postcondición", "Usuario registrado en la base de datos y redirigido a la página de inicio de sesión."),
    ])
    body(doc, "Secuencia Normal:", bold=True)
    make_table(doc, [
        ["Paso", "Acción"],
        ["1", "El usuario navega a la página de registro (/register)."],
        ["2", "El sistema muestra el formulario con campos: nombre completo, correo, contraseña."],
        ["3", "El usuario ingresa sus datos y envía el formulario."],
        ["4", "El cliente valida los campos requeridos usando JavaScript."],
        ["5", "Se envía petición POST asíncrona (AJAX/Fetch) al controlador AuthController."],
        ["6", "El controlador sanitiza los datos y ejecuta sp_register_user en MySQL."],
        ["7", "La contraseña se hashea con bcrypt antes del almacenamiento."],
        ["8", "El sistema retorna JSON con confirmación o errores de validación."],
    ], [0.8, 5.7])
    body(doc, "Excepciones:", bold=True)
    make_table(doc, [
        ["Paso", "Acción"],
        ["1", "Si el correo ya existe, el sistema retorna error 'El correo ya está registrado'."],
        ["2", "Si hay campos vacíos, se muestran mensajes de validación."],
    ], [0.8, 5.7])
    body(doc, "Rendimiento: Cota de tiempo: 2 segundos. Frecuencia esperada: 50 veces por semana. Importancia: Muy importante. Urgencia: Inmediatamente.")

    # CU-02
    heading(doc, "7.1.2 Caso de Uso — Inicio de Sesión", 3)
    make_kv_table(doc, [
        ("Identificador", "CU-02"),
        ("Nombre", "Inicio de Sesión"),
        ("Actor Principal", "Estudiante / Usuario Regular / Administrador"),
        ("Descripción", "Autentica a un usuario con sus credenciales (correo y contraseña)."),
        ("Precondición", "Tener una cuenta registrada y activa en el sistema."),
        ("Postcondición", "Sesión iniciada, acceso al feed de publicaciones (user) o dashboard (admin)."),
    ])
    body(doc, "Secuencia Normal:", bold=True)
    make_table(doc, [
        ["Paso", "Acción"],
        ["1", "El usuario navega a la página de inicio de sesión (/login)."],
        ["2", "El sistema muestra formulario con campos: correo, contraseña."],
        ["3", "El usuario ingresa credenciales y envía el formulario."],
        ["4", "El cliente valida campos requeridos con JavaScript."],
        ["5", "Se envía petición POST al AuthController."],
        ["6", "El controlador ejecuta sp_login_user para verificar existencia del correo."],
        ["7", "Se verifica la contraseña con password_verify() contra el hash bcrypt."],
        ["8", "El sistema regenera el ID de sesión y establece variables de sesión."],
        ["9", "El usuario es redirigido según su rol (user → /posts, admin → /dashboard)."],
    ], [0.8, 5.7])
    body(doc, "Excepciones:", bold=True)
    make_table(doc, [
        ["Paso", "Acción"],
        ["1", "Si el correo no existe: mensaje 'Correo no encontrado'."],
        ["2", "Si la contraseña es incorrecta: mensaje 'Contraseña incorrecta'."],
        ["3", "Si la cuenta está inactiva: mensaje 'Cuenta desactivada'."],
    ], [0.8, 5.7])
    body(doc, "Rendimiento: Cota de tiempo: 2 segundos. Frecuencia esperada: 200 veces por día. Importancia: Muy importante. Urgencia: Inmediatamente.")

    # CU-03
    heading(doc, "7.1.3 Caso de Uso — Crear Publicación", 3)
    make_kv_table(doc, [
        ("Identificador", "CU-03"),
        ("Nombre", "Crear Publicación"),
        ("Actor Principal", "Estudiante / Usuario Regular"),
        ("Descripción", "El usuario redacta un texto o sube una imagen al feed general de la red social."),
        ("Precondición", "Estar autenticado en el sistema."),
        ("Postcondición", "Publicación guardada en la base de datos y visible para amigos en el feed."),
    ])
    body(doc, "Secuencia Normal:", bold=True)
    make_table(doc, [
        ["Paso", "Acción"],
        ["1", "El usuario navega a la sección 'Agregar Post' (/addPost)."],
        ["2", "El sistema muestra formulario con campos: contenido (textarea) e imagen (opcional)."],
        ["3", "El usuario ingresa el texto y opcionalmente selecciona una imagen."],
        ["4", "El cliente valida tipos de archivo (JPEG, PNG, GIF) y tamaño (máx. 5MB)."],
        ["5", "Se envía petición POST con FormData al PostController."],
        ["6", "El controlador sanitiza el texto, procesa la imagen y ejecuta sp_create_post."],
        ["7", "La imagen se almacena en assets/imagesPosts/ con nombre único."],
        ["8", "El sistema retorna JSON con confirmación y redirige al feed."],
    ], [0.8, 5.7])
    body(doc, "Excepciones:", bold=True)
    make_table(doc, [
        ["Paso", "Acción"],
        ["1", "Si el archivo no es imagen válida: mensaje de error de formato."],
        ["2", "Si el archivo excede 5MB: mensaje de error de tamaño."],
        ["3", "Si el contenido está vacío y no hay imagen: mensaje de validación."],
    ], [0.8, 5.7])
    body(doc, "Rendimiento: Cota de tiempo: 3 segundos. Frecuencia esperada: 100 veces por día. Importancia: Muy importante. Urgencia: Inmediatamente.")

    # CU-04 through CU-10 summarized
    heading(doc, "7.1.4 — 7.1.10 Casos de Uso Adicionales", 3)
    additional_cus = [
        ("CU-04 — Comentar Publicación", "El usuario agrega un comentario en una publicación existente. Precondición: publicación existente y sesión iniciada. Postcondición: comentario visible con paginación."),
        ("CU-05 — Dar Me Gusta", "El usuario reacciona positivamente a una publicación, comentario o respuesta. Toggle like/unlike con actualización de contador en tiempo real."),
        ("CU-06 — Enviar Solicitud de Amistad", "El usuario busca otro perfil y envía una solicitud para conectar. Estado 'pending' en tabla friend_requests."),
        ("CU-07 — Aceptar/Rechazar Amistad", "El receptor gestiona solicitudes entrantes. Si acepta, ambos usuarios se registran en tabla friends."),
        ("CU-08 — Editar Perfil", "El usuario modifica su nombre completo, biografía o foto de perfil. Se genera registro de auditoría automático."),
        ("CU-09 — Eliminar Publicación", "El autor elimina su publicación mediante soft delete (active=0). Verificación de ownership."),
        ("CU-10 — Búsqueda de Usuarios", "Motor de búsqueda por nombre con resultados renderizados como tarjetas de usuario con opciones de amistad."),
    ]
    for title, desc in additional_cus:
        body(doc, f"{title}: {desc}")

    doc.add_paragraph()
    heading(doc, "7.2 Actores del Sistema", 2)

    heading(doc, "Actor #1: Estudiante / Usuario Regular", 3)
    make_kv_table(doc, [
        ("Actor", "Estudiante / Usuario Regular"),
        ("Tipo", "Primario"),
        ("Casos de Uso", "CU-01 a CU-10: Registro, Login, Publicaciones, Comentarios, Likes, Amistades, Perfil, Búsqueda."),
        ("Descripción", "El Usuario Regular es el estudiante universitario que utiliza el sistema para conectarse con la comunidad, compartir publicaciones, interactuar mediante comentarios y likes, gestionar su red de amistades y personalizar su perfil. Es el actor principal y más frecuente del sistema."),
    ])

    heading(doc, "Actor #2: Administrador", 3)
    make_kv_table(doc, [
        ("Actor", "Administrador"),
        ("Tipo", "Primario"),
        ("Casos de Uso", "Login (rol admin), Visualizar Dashboard, Consultar Estadísticas y Gráficas, Generar Reporte PDF, Ver Rankings."),
        ("Descripción", "El Administrador tiene credenciales especiales (admin@gmail.com) que le otorgan acceso al dashboard administrativo. Desde allí puede visualizar métricas generales de la plataforma, monitorear la actividad de los usuarios, analizar gráficas de tendencias, consultar rankings y generar reportes PDF."),
    ])

    page_break(doc)


def build_design(doc):
    """Section 8: Unidad de Diseño"""
    heading(doc, "8. UNIDAD DE DISEÑO", 1)
    heading(doc, "8.1 Diseño de Interfaces", 2)
    body(doc, "El enfoque de diseño de UNIRED está altamente centrado en la usabilidad y la estética moderna. A continuación, se describen los elementos visuales y decisiones de diseño clave:")

    design_items = [
        ("Tipografía", "Roboto (Google Fonts) en múltiples pesos (300, 400, 700) para garantizar jerarquía visual. Los títulos usan peso 700, el cuerpo 400 y textos secundarios 300."),
        ("Paleta de Colores", "Color primario: Azul Turquesa (#4db8c4). Fondos: blanco (#ffffff) y gris claro (#f4f6f9). Texto principal: gris oscuro (#333333). Alertas de error: rojo suave, confirmaciones: verde."),
        ("Layout Desktop", "Sidebar fija a la izquierda (250px) con navegación principal y foto de perfil. Feed central scrolleable con ancho máximo de 680px. Dashboard admin con widgets en grid adaptable."),
        ("Layout Móvil", "La sidebar se transforma en un drawer oculto que se despliega al interactuar con el botón de menú (hamburguesa). Las publicaciones ocupan el ancho completo. Tablas scrollables horizontalmente."),
        ("Sidebar", "Componente reutilizable (views/assets/sidebar.php) incluido en todas las vistas autenticadas. Contiene: foto de perfil circular, nombre del usuario, enlaces de navegación con iconos SVG."),
        ("Publicaciones", "Tarjetas blancas con sombra sutil. Cabecera con foto y nombre del autor + timestamp. Cuerpo con texto y/o imagen. Pie con botones de Like, Comentarios y menú de opciones."),
        ("Comentarios", "Sección colapsable debajo de cada publicación. Muestra primeros 3 comentarios con botón 'Cargar más'. Cada comentario muestra foto del autor, nombre con enlace, texto y timestamp."),
        ("Modales", "Modal de confirmación para acciones críticas (eliminar cuenta, publicación, amigo) con overlay oscuro. Modal de pantalla completa para visualización de imágenes."),
        ("Dashboard Admin", "5 tarjetas de resumen (KPI cards) con iconos y valores numéricos. Gráficas Chart.js: líneas, barras, dona. 4 tablas de rankings."),
        ("Formularios", "Campos con borde sutil, enfoque con resaltado azul turquesa. Validación en tiempo real con JavaScript. Mensajes de error en rojo. Botones con transiciones 0.3s ease-in-out."),
        ("Microinteracciones", "Transiciones 0.3s en botones y tarjetas. Esqueletos de carga durante peticiones asíncronas. Animación de corazón al dar like. Toast de notificaciones."),
    ]
    for title, desc in design_items:
        body(doc, f"{title}:", bold=True)
        body(doc, desc)

    body(doc, "Todas las interfaces fueron diseñadas siguiendo el enfoque Mobile First, garantizando una experiencia de usuario óptima en dispositivos móviles (desde 320px), tablets (768px) y escritorio (1024px+).")
    page_break(doc)


def build_database(doc):
    """Section 9: Base de Datos"""
    heading(doc, "9. BASE DE DATOS", 1)
    body(doc, "La estructura de base de datos relacional de UNIRED está optimizada para la integridad referencial y las consultas veloces mediante Stored Procedures. Se utiliza el motor InnoDB de MySQL con codificación utf8mb4_unicode_ci para soporte completo de caracteres.")

    heading(doc, "9.1 Tablas Principales", 2)
    make_table(doc, [
        ["Tabla", "Descripción"],
        ["users", "Usuarios registrados. PK: user_id. Campos: full_name, biography, profile_picture, email (UNIQUE), password (bcrypt), registration_date, role (user/admin), active, updated_at."],
        ["posts", "Publicaciones. PK: post_id. FK: user_id. Campos: content, image, created_at, updated_at, active (soft delete)."],
        ["comments", "Comentarios. PK: comment_id. FK: post_id, user_id. Campos: content, created_at, active."],
        ["replies", "Respuestas anidadas. PK: reply_id. FK: comment_id, user_id. Campos: content, created_at, active."],
        ["likes", "Likes en publicaciones. PK: like_id. FK: post_id, user_id. UNIQUE(post_id, user_id)."],
        ["comment_likes", "Likes en comentarios. FK: comment_id, user_id. UNIQUE(comment_id, user_id)."],
        ["reply_likes", "Likes en respuestas. FK: reply_id, user_id. UNIQUE(reply_id, user_id)."],
        ["friend_requests", "Solicitudes de amistad. PK: request_id. FK: sender_id, receiver_id. Campos: status, request_date, response_date."],
        ["friends", "Amistades confirmadas. PK: friendship_id. FK: user_id1, user_id2. UNIQUE(user_id1, user_id2)."],
        ["hidden_comments", "Comentarios ocultos. FK: user_id, comment_id. UNIQUE(user_id, comment_id)."],
        ["user_update_log", "Auditoría de perfil. PK: log_id. FK: user_id. Campos: old_full_name, new_full_name, old_biography, new_biography, change_date."],
    ], [1.3, 5.2])

    heading(doc, "9.2 Procedimientos Almacenados", 2)
    body(doc, "El sistema utiliza 26 Stored Procedures para encapsular la lógica de negocio en la capa de base de datos:")
    make_table(doc, [
        ["Procedimiento", "Descripción"],
        ["sp_register_user", "Registra un nuevo usuario con validación de correo duplicado."],
        ["sp_login_user", "Recupera datos de usuario por correo para autenticación."],
        ["sp_create_post", "Crea una nueva publicación y retorna su ID."],
        ["sp_create_comment", "Crea un comentario y retorna su ID."],
        ["sp_get_comments_by_post", "Obtiene comentarios de una publicación con datos del autor."],
        ["sp_delete_comment", "Soft delete de un comentario (active=0)."],
        ["sp_get_comment_count", "Conteo de comentarios activos de un post."],
        ["sp_get_comment_by_id", "Obtiene un comentario específico por ID."],
        ["sp_create_reply", "Crea una respuesta a un comentario."],
        ["sp_get_replies_by_comment", "Obtiene respuestas de un comentario con datos del autor."],
        ["sp_delete_reply", "Soft delete de una respuesta."],
        ["sp_get_reply_count", "Conteo de respuestas activas de un comentario."],
        ["sp_get_reply_by_id", "Obtiene una respuesta específica por ID."],
        ["sp_add_like", "Agrega un like a una publicación (INSERT IGNORE)."],
        ["sp_remove_like", "Elimina un like de una publicación."],
        ["sp_get_like_count", "Conteo de likes de una publicación."],
        ["sp_has_liked", "Verifica si un usuario dio like a una publicación."],
        ["sp_get_user_likes", "Obtiene todos los likes de un usuario."],
        ["sp_get_post_likers", "Obtiene usuarios que dieron like a un post."],
        ["sp_add_comment_like", "Agrega like a un comentario."],
        ["sp_remove_comment_like", "Elimina like de comentario."],
        ["sp_get_comment_like_count", "Conteo de likes de un comentario."],
        ["sp_has_comment_liked", "Verifica like en comentario."],
        ["sp_add_reply_like", "Agrega like a una respuesta."],
        ["sp_remove_reply_like", "Elimina like de respuesta."],
        ["sp_get_reply_like_count", "Conteo de likes de una respuesta."],
        ["sp_has_reply_liked", "Verifica like en respuesta."],
    ], [2.4, 4.1])

    heading(doc, "9.3 Vistas y Triggers", 2)
    body(doc, "Vistas:", bold=True)
    body(doc, "v_posts_stats: Vista que combina la tabla posts con users, conteo de likes y conteo de comentarios activos para renderizar eficientemente el feed principal.")
    body(doc, "Triggers:", bold=True)
    body(doc, "trg_user_update_log: Trigger BEFORE UPDATE en la tabla users que registra automáticamente los cambios de nombre y biografía en la tabla user_update_log para fines de auditoría y trazabilidad.")
    page_break(doc)


def build_sequence_diagrams(doc):
    """Section 10: Diagramas de Secuencia"""
    heading(doc, "10. DIAGRAMAS DE SECUENCIA", 1)
    body(doc, "A continuación se describen los diagramas de secuencia de los principales flujos del sistema, siguiendo el patrón de comunicación entre el Cliente (Frontend JavaScript), el Enrutador PHP, el Controlador, el Modelo y la Base de Datos MySQL.")

    diagrams = [
        ("10.1 Registro de Usuario", [
            "1. Usuario accede a /register",
            "2. Sistema muestra formulario de registro",
            "3. Usuario completa datos y envía",
            "4. Cliente (JS) valida campos requeridos",
            "5. Fetch POST /register con datos JSON",
            "6. Router → AuthController::register()",
            "7. Controller sanitiza inputs",
            "8. Controller → UserModel::registerUser()",
            "9. Model ejecuta sp_register_user",
            "10. MySQL valida y registra usuario",
            "11. Model retorna resultado",
            "12. Controller retorna JSON response",
            "13. Cliente muestra confirmación y redirige a /login",
        ]),
        ("10.2 Crear Publicación", [
            "1. Usuario accede a /addPost",
            "2. Sistema muestra formulario de creación",
            "3. Usuario ingresa texto y/o selecciona imagen",
            "4. Cliente valida tipo de archivo y tamaño",
            "5. Fetch POST /posts con FormData",
            "6. Router → PostController::store()",
            "7. Controller sanitiza texto",
            "8. Controller procesa upload de imagen",
            "9. Controller → PostModel::createPost()",
            "10. Model ejecuta sp_create_post",
            "11. MySQL inserta registro en tabla posts",
            "12. Model retorna post_id",
            "13. Controller retorna JSON success",
            "14. Cliente redirige a /posts (feed)",
        ]),
        ("10.3 Dar Me Gusta a Publicación", [
            "1. Usuario hace clic en icono de corazón en un post",
            "2. Cliente (JS) captura evento click",
            "3. Cliente verifica estado actual (liked/unliked)",
            "4. Si no ha dado like: Fetch POST /posts/:id/like",
            "5. Si ya dio like: Fetch DELETE /posts/:id/like",
            "6. Router → PostController::like() / unlike()",
            "7. Controller verifica autenticación",
            "8. Controller → LikeModel::addLike() / removeLike()",
            "9. Model ejecuta sp_add_like / sp_remove_like",
            "10. MySQL INSERT IGNORE o DELETE",
            "11. Model retorna affected_rows",
            "12. Controller retorna JSON con nuevo conteo",
            "13. Cliente actualiza contador e icono sin recargar",
        ]),
        ("10.4 Enviar Solicitud de Amistad", [
            "1. Usuario busca perfil → clic en 'Enviar Solicitud'",
            "2. Cliente (JS) captura evento",
            "3. Fetch POST /friend/request/:id",
            "4. Router → FriendController::sendRequestById()",
            "5. Controller verifica autenticación",
            "6. Controller → FriendModel::sendRequest()",
            "7. Model valida que no exista solicitud previa",
            "8. Model INSERT en friend_requests (status='pending')",
            "9. Model retorna resultado",
            "10. Controller retorna JSON status",
            "11. Cliente actualiza botón a 'Solicitud Enviada'",
        ]),
        ("10.5 Dashboard Administrativo", [
            "1. Admin inicia sesión → redirigido a /dashboard",
            "2. Router verifica rol 'admin' → requireAdmin()",
            "3. Sistema carga vista admin/dashboard.php",
            "4. Cliente (JS) ejecuta múltiples Fetch en paralelo",
            "5. GET /admin/stats/summary → AdminController",
            "6. AdminModel::getSummaryStats() consulta BD",
            "7. GET /admin/stats/activity-timeline",
            "8. GET /admin/stats/engagement-breakdown",
            "9. GET /admin/stats/user-growth",
            "10. Cliente renderiza gráficas Chart.js",
            "11. Cliente llena tablas de rankings",
        ]),
        ("10.6 Generar Reporte PDF", [
            "1. Admin hace clic en 'Descargar Reporte PDF'",
            "2. Cliente abre /admin/stats/pdf",
            "3. PdfController::downloadStatsPdf()",
            "4. AdminModel::getUsersWithMostPosts()",
            "5. AdminModel::getUsersWithMostFriends()",
            "6. AdminModel::getPostsWithMostComments()",
            "7. AdminModel::getPostsWithMostLikes()",
            "8. Controller instancia TCPDF y construye documento",
            "9. Headers Content-Type: application/pdf",
            "10. TCPDF genera PDF con tablas y datos",
            "11. Navegador recibe PDF y lo descarga",
        ]),
    ]

    for title, steps in diagrams:
        heading(doc, title, 2)
        for step in steps:
            bullet(doc, step)

    page_break(doc)


def build_class_diagrams(doc):
    """Section 11: Diagramas de Clases"""
    heading(doc, "11. DIAGRAMAS DE CLASES", 1)
    body(doc, "Los diagramas de clases representan la estructura estática del sistema UNIRED, modelando las entidades principales y sus relaciones.")

    heading(doc, "11.1 Estructura General MVC", 2)
    body(doc, "El sistema sigue el patrón de arquitectura MVC (Modelo-Vista-Controlador) con las siguientes clases principales:")

    body(doc, "Capa de Controladores (app/Controllers/):", bold=True)
    bullet(doc, "AuthController: Gestiona registro, inicio de sesión y cierre de sesión. Métodos: register(), login(), logout().")
    bullet(doc, "PostController: CRUD de publicaciones, likes y comentarios. Métodos: store(), update(), destroy(), like(), unlike(), addComment(), deleteComment(), getComments(), addReply(), deleteReply(), getReplies(), likeComment(), likeReply().")
    bullet(doc, "UserController: Gestión de perfiles de usuario. Métodos: show(), edit(), update(), deleteAccount().")
    bullet(doc, "FriendController: Gestión de relaciones de amistad con 10+ métodos para enviar, aceptar, rechazar, cancelar solicitudes y gestionar amigos.")
    bullet(doc, "AdminController: Endpoints de estadísticas. 13 métodos incluyendo getSummaryStats(), getActivityTimeline(), getEngagementBreakdown(), getUserGrowth(), getPostsByDayOfWeek(), etc.")
    bullet(doc, "PdfController: Generación de reportes PDF. Método: downloadStatsPdf().")

    body(doc, "Capa de Modelos (app/Models/):", bold=True)
    bullet(doc, "UserModel: Operaciones CRUD sobre tabla users. Utiliza sp_register_user y sp_login_user.")
    bullet(doc, "PostModel: Operaciones CRUD sobre tabla posts.")
    bullet(doc, "CommentModel: Operaciones sobre tabla comments con soft delete.")
    bullet(doc, "ReplyModel: Operaciones sobre tabla replies con soft delete.")
    bullet(doc, "LikeModel: Gestión de likes en publicaciones con toggle.")
    bullet(doc, "CommentLikeModel: Gestión de likes en comentarios.")
    bullet(doc, "ReplyLikeModel: Gestión de likes en respuestas.")
    bullet(doc, "FriendModel: Gestión de amistades con validación de estados.")
    bullet(doc, "AdminModel: Consultas estadísticas agregadas para dashboard.")

    body(doc, "Capa de Componentes (app/Components/):", bold=True)
    bullet(doc, "Alert: Renderiza mensajes flash (éxito, error, info).")
    bullet(doc, "Post: Renderiza una publicación completa con comentarios y botones.")
    bullet(doc, "FriendCard: Tarjeta de amigo/sugerencia con botones contextuales.")
    bullet(doc, "Profile: Cabecera de perfil con estadísticas, foto y acciones.")

    heading(doc, "11.2 Relaciones entre Clases", 2)
    make_table(doc, [
        ["Relación", "Descripción"],
        ["Usuario → Publicación (1:N)", "Un usuario puede crear múltiples publicaciones. FK user_id en posts."],
        ["Publicación → Comentario (1:N)", "Una publicación puede tener múltiples comentarios. FK post_id en comments."],
        ["Comentario → Respuesta (1:N)", "Un comentario puede tener múltiples respuestas. FK comment_id en replies."],
        ["Usuario → Like (N:M)", "Relación muchos a muchos entre usuarios y publicaciones vía tabla likes. UNIQUE(post_id, user_id)."],
        ["Usuario → Amigo (N:M)", "Relación muchos a muchos entre usuarios vía tabla friends. UNIQUE(user_id1, user_id2)."],
        ["Usuario → Solicitud (N:M)", "Relación bidireccional vía friend_requests con estados (pending/accepted/rejected/cancelled)."],
        ["Usuario → Auditoría (1:N)", "Trigger BEFORE UPDATE registra cambios de perfil en user_update_log."],
    ], [2.5, 4.0])
    page_break(doc)


def build_testing(doc):
    """Section 12: Pruebas"""
    heading(doc, "12. PRUEBAS", 1)
    body(doc, "Para garantizar la calidad del software, se realizaron pruebas de caja blanca y caja negra, documentando los casos de prueba y sus resultados.")

    heading(doc, "12.1 Pruebas de Caja Blanca", 2)
    body(doc, "Las pruebas de caja blanca se enfocaron en verificar la lógica interna del código, específicamente en los Stored Procedures y controladores.")

    heading(doc, "12.1.1 Caso de Prueba CB-01: Registro de Usuario", 3)
    make_kv_table(doc, [
        ("ID Caso de Prueba", "CB-01"),
        ("Módulo", "Autenticación — Registro de Usuario"),
        ("Objetivo", "Verificar que sp_register_user detecte correctamente correos duplicados y registre usuarios válidos."),
        ("Entradas", "Caso 1: correo nuevo. Caso 2: correo existente (admin@gmail.com)."),
        ("Resultado Esperado", "Caso 1: usuario registrado. Caso 2: error 'El correo ya está registrado'."),
        ("Resultado Obtenido", "Caso 1: registro exitoso. Caso 2: SIGNAL SQLSTATE 45000 lanzada correctamente."),
        ("Estado", "APROBADO ✓"),
    ])

    heading(doc, "12.1.2 Caso de Prueba CB-02: Sistema de Likes", 3)
    make_kv_table(doc, [
        ("ID Caso de Prueba", "CB-02"),
        ("Módulo", "Publicaciones — Sistema de Likes"),
        ("Objetivo", "Verificar que sp_add_like use INSERT IGNORE correctamente y sp_remove_like elimine el registro."),
        ("Entradas", "post_id=1, user_id=2. Ejecutar sp_add_like dos veces."),
        ("Resultado Esperado", "Primera: 1 fila afectada. Segunda: 0 filas (ignorada)."),
        ("Resultado Obtenido", "Primera: affected_rows=1. Segunda: affected_rows=0. Sin duplicados."),
        ("Estado", "APROBADO ✓"),
    ])

    heading(doc, "12.2 Pruebas de Caja Negra", 2)
    body(doc, "Las pruebas de caja negra verifican el comportamiento del sistema desde la perspectiva del usuario final.")

    heading(doc, "12.2.1 Pruebas de Aceptación", 3)
    body(doc, "Se verificaron los criterios de aceptación definidos para cada requerimiento funcional:")
    make_table(doc, [
        ["Requerimiento", "Criterio", "Resultado"],
        ["RF-01 Registro", "Usuario puede registrarse con datos válidos", "APROBADO ✓"],
        ["RF-01 Registro", "Sistema rechaza correo duplicado", "APROBADO ✓"],
        ["RF-02 Login", "Usuario puede iniciar sesión con credenciales correctas", "APROBADO ✓"],
        ["RF-02 Login", "Sistema rechaza credenciales incorrectas", "APROBADO ✓"],
        ["RF-03 Publicación", "Usuario puede crear publicación con texto e imagen", "APROBADO ✓"],
        ["RF-03 Publicación", "Sistema rechaza archivos no permitidos", "APROBADO ✓"],
        ["RF-04 Comentarios", "Comentarios se cargan con paginación (3 en 3)", "APROBADO ✓"],
        ["RF-05 Amigos", "Flujo completo enviar-aceptar-rechazar funciona", "APROBADO ✓"],
        ["RF-06 Perfil", "Usuario puede editar nombre, bio y foto", "APROBADO ✓"],
        ["RF-07 Dashboard", "Admin ve estadísticas y gráficas correctamente", "APROBADO ✓"],
        ["RF-08 PDF", "Reporte PDF se genera y descarga correctamente", "APROBADO ✓"],
        ["RF-09 Eliminar Post", "Solo el autor puede eliminar su publicación", "APROBADO ✓"],
        ["RF-15 Eliminar Cuenta", "Eliminación en cascada de todos los datos", "APROBADO ✓"],
    ], [1.7, 3.0, 1.8])

    heading(doc, "12.2.2 Pruebas de Usabilidad de Interfaz", 3)
    bullet(doc, "Diseño Responsive: Verificado en resoluciones 320px (móvil), 768px (tablet) y 1920px (desktop). Sidebar funciona como drawer en móvil.")
    bullet(doc, "Navegación: Todos los enlaces de la sidebar dirigen a las vistas correctas. Cerrar sesión limpia la sesión adecuadamente.")
    bullet(doc, "Modales de Confirmación: Acciones críticas muestran modal de confirmación antes de ejecutarse.")
    bullet(doc, "Feedback Visual: Botones cambian de estado al hover. Operaciones asíncronas muestran indicadores de carga. Mensajes flash tras operaciones.")
    bullet(doc, "Accesibilidad: Contraste adecuado entre texto y fondo. Tamaños de fuente legibles. Formularios con etiquetas claras.")
    body(doc, "Resultado: APROBADO ✓ — Todas las pruebas de usabilidad fueron superadas satisfactoriamente.", bold=True)

    heading(doc, "12.2.3 Pruebas de Seguridad", 3)
    bullet(doc, "Inyección SQL: Verificado que todas las consultas usan PDO con prepared statements. No se pudo inyectar SQL malicioso.")
    bullet(doc, "XSS (Cross-Site Scripting): Todas las salidas pasan por htmlspecialchars() o safe_output(). Scripts maliciosos no se ejecutan.")
    bullet(doc, "CSRF: Operaciones de cambio de estado (POST, PUT, DELETE) requieren sesión activa.")
    bullet(doc, "Control de Acceso: Usuarios sin autenticación no pueden acceder a rutas protegidas. Usuarios regulares no acceden al dashboard.")
    bullet(doc, "Subida de Archivos: Solo se aceptan imágenes (JPEG, PNG, GIF). Tamaño máximo 5 MB. Archivos PHP renombrados no se ejecutan.")
    body(doc, "Resultado: APROBADO ✓ — El sistema pasó todas las pruebas de seguridad sin vulnerabilidades detectadas.", bold=True)
    page_break(doc)


def build_deployment(doc):
    """Section 13: Plan de Implantación"""
    heading(doc, "13. PLAN DE IMPLANTACIÓN", 1)
    body(doc, "El plan de implantación considera las 5 acciones necesarias para poner el sistema en funcionamiento en el entorno de producción.")

    actions = [
        ("13.1 Instalación de Software", [
            "1. Configurar servidor web Apache con mod_rewrite habilitado.",
            "2. Instalar PHP 7.4+ con extensiones: PDO, MySQL, mbstring, fileinfo, gd.",
            "3. Instalar Composer para gestión de dependencias.",
            "4. Clonar repositorio del proyecto en el directorio raíz del servidor web.",
            "5. Ejecutar 'composer install' para instalar dependencias (vlucas/phpdotenv, TCPDF).",
            "6. Configurar archivo .env con credenciales de base de datos.",
            "7. Configurar permisos de escritura en directorios assets/imagesProfile/ y assets/imagesPosts/.",
        ]),
        ("13.2 Instalación de Hardware y Telecomunicaciones", [
            "1. Contratar servicio de hosting en la nube (Hostinger).",
            "2. Configurar dominio (darksalmon-jellyfish-884197.hostingersite.com).",
            "3. Configurar certificado SSL para HTTPS.",
            "4. Verificar conectividad de red y ancho de banda adecuado.",
            "5. Configurar backups automáticos de base de datos (diarios).",
        ]),
        ("13.3 Carga de Datos", [
            "1. Ejecutar script database/unired_db.sql para crear estructura de base de datos.",
            "2. Verificar creación de 11 tablas, 26 Stored Procedures, 1 Trigger y 1 Vista.",
            "3. Ejecutar database/seed.php para poblar datos de prueba (25 usuarios, 55 posts, 150 comentarios, 80 replies, 300 likes, 100 comment likes, 60 reply likes, 40 amistades).",
            "4. Verificar integridad de datos con consultas de validación.",
            "5. Crear cuenta de administrador: admin@gmail.com / Admin123?.",
        ]),
        ("13.4 Capacitación", [
            "1. Sesión para administradores (1 hora): acceso al dashboard, interpretación de gráficas, generación de reportes PDF.",
            "2. Sesión para usuarios finales (1 hora): registro, login, publicaciones, comentarios, gestión de amistades, edición de perfil.",
            "3. Entrega de manual de usuario digital.",
            "4. Video tutorial de uso del sistema.",
        ]),
        ("13.5 Conversión", [
            "1. Estrategia de adopción gradual: Fase 1 (grupo piloto 50 estudiantes), Fase 2 (toda la comunidad), Fase 3 (operación normal con monitoreo).",
            "2. Período de operación paralela opcional (no aplica — no hay sistema previo).",
            "3. Monitoreo de métricas de adopción: nuevos registros, publicaciones, interacciones.",
            "4. Canal de soporte para reportar bugs o sugerencias.",
            "5. Plan de mantenimiento: actualizaciones mensuales de seguridad y dependencias.",
        ]),
    ]

    for title, items in actions:
        heading(doc, title, 2)
        for item in items:
            bullet(doc, item)

    page_break(doc)


def build_user_manual(doc):
    """Section 14: Manual de Usuario"""
    heading(doc, "14. MANUAL DE USUARIO", 1)
    body(doc, "Este manual describe los pasos necesarios para utilizar el Sistema Integral de Red Social Universitaria UNIRED.")

    heading(doc, "14.1 Instalación", 2)
    body(doc, "UNIRED es una aplicación web. Los usuarios finales no requieren instalar ningún software adicional. Solo necesitan un navegador web actualizado (Chrome, Firefox, Safari, Edge) y conexión a internet. La URL de acceso es: https://darksalmon-jellyfish-884197.hostingersite.com/")

    heading(doc, "14.2 Uso del Sistema", 2)

    heading(doc, "14.2.1 Registro de Usuario", 3)
    bullet(doc, "1. Acceder a la URL del sistema.")
    bullet(doc, "2. Hacer clic en 'Registrarse' o navegar a /register.")
    bullet(doc, "3. Completar el formulario con: nombre completo, correo electrónico y contraseña.")
    bullet(doc, "4. Hacer clic en 'Registrarse'.")
    bullet(doc, "5. Si el registro es exitoso, será redirigido a la página de inicio de sesión.")

    heading(doc, "14.2.2 Inicio de Sesión", 3)
    bullet(doc, "1. Navegar a /login.")
    bullet(doc, "2. Ingresar correo electrónico y contraseña.")
    bullet(doc, "3. Hacer clic en 'Iniciar Sesión'.")
    bullet(doc, "4. El sistema redirigirá al feed de publicaciones (rol user) o al dashboard (rol admin).")

    heading(doc, "14.2.3 Navegación Principal", 3)
    bullet(doc, "Inicio (Feed): Muestra las publicaciones de todos los usuarios.")
    bullet(doc, "Amigos: Gestiona lista de amigos, solicitudes pendientes y búsqueda de usuarios.")
    bullet(doc, "Mi Perfil: Visualiza tu perfil y tus publicaciones.")
    bullet(doc, "Agregar Post: Crea una nueva publicación con texto e imagen opcional.")
    bullet(doc, "Editar Perfil: Modifica tu nombre, biografía y foto de perfil.")
    bullet(doc, "Cerrar Sesión: Finaliza la sesión actual.")
    body(doc, "En dispositivos móviles, la sidebar se oculta y se muestra al presionar el botón de menú (☰).")

    heading(doc, "14.2.4 Crear una Publicación", 3)
    bullet(doc, "1. Hacer clic en 'Agregar Post' en la sidebar.")
    bullet(doc, "2. Escribir el contenido en el campo de texto.")
    bullet(doc, "3. Opcionalmente, seleccionar una imagen (JPEG, PNG o GIF, máximo 5 MB).")
    bullet(doc, "4. Hacer clic en 'Publicar'.")
    bullet(doc, "5. La publicación aparecerá en el feed general.")

    heading(doc, "14.2.5 Interactuar con Publicaciones", 3)
    bullet(doc, "Me Gusta: Haz clic en el icono de corazón para dar like. Vuelve a hacer clic para quitarlo.")
    bullet(doc, "Comentarios: Haz clic en 'Comentarios' para ver los existentes. Escribe y presiona Enter para comentar.")
    bullet(doc, "Responder: Dentro de un comentario, haz clic en 'Responder' para agregar una respuesta.")
    bullet(doc, "Ver Imagen: Haz clic en una imagen para verla en pantalla completa. Haz clic en X para cerrar.")
    bullet(doc, "Editar/Eliminar: Si eres el autor, verás un menú de opciones (···) para editar o eliminar.")

    heading(doc, "14.2.6 Gestionar Amigos", 3)
    bullet(doc, "1. Navegar a la sección 'Amigos'.")
    bullet(doc, "2. Usar la barra de búsqueda para encontrar usuarios por nombre.")
    bullet(doc, "3. Hacer clic en 'Enviar Solicitud' para conectar con otro usuario.")
    bullet(doc, "4. En la pestaña 'Solicitudes', ver las solicitudes entrantes y aceptarlas o rechazarlas.")
    bullet(doc, "5. En la pestaña 'Mis Amigos', ver la lista de amigos y opción para eliminar.")

    heading(doc, "14.2.7 Editar Perfil", 3)
    bullet(doc, "1. Hacer clic en 'Editar Perfil' en la sidebar.")
    bullet(doc, "2. Modificar nombre completo y/o biografía.")
    bullet(doc, "3. Para cambiar la foto, hacer clic en 'Seleccionar archivo' y elegir una imagen.")
    bullet(doc, "4. Hacer clic en 'Guardar Cambios'.")
    bullet(doc, "5. Los cambios se reflejarán inmediatamente en el perfil.")

    heading(doc, "14.2.8 Dashboard Administrativo", 3)
    bullet(doc, "Disponible solo para usuarios con rol 'admin' (admin@gmail.com).")
    bullet(doc, "Ver tarjetas de resumen con totales de usuarios, publicaciones, comentarios, likes y amistades.")
    bullet(doc, "Explorar gráficas interactivas de actividad, engagement, crecimiento de usuarios, etc.")
    bullet(doc, "Consultar rankings de usuarios con más publicaciones y amigos.")
    bullet(doc, "Consultar rankings de publicaciones con más comentarios y likes.")
    bullet(doc, "Hacer clic en 'Descargar Reporte PDF' para obtener un documento con estadísticas.")

    heading(doc, "14.3 Solución de Problemas", 2)
    make_table(doc, [
        ["Problema", "Solución"],
        ["No puedo iniciar sesión", "Verificar que el correo y contraseña sean correctos. Contactar al administrador si olvidó la contraseña."],
        ["No puedo subir una imagen", "Verificar que la imagen sea JPEG, PNG o GIF y no exceda 5 MB."],
        ["No veo mis publicaciones", "Verificar que ha iniciado sesión. Si persiste, cerrar sesión y volver a iniciar."],
        ["El sistema está lento", "Verificar la conexión a internet. Contactar al administrador si el problema persiste."],
        ["Error al eliminar cuenta", "Asegurarse de haber iniciado sesión. La eliminación es permanente."],
        ["La sidebar no se muestra en móvil", "Presionar el botón de menú (☰) en la esquina superior para desplegar la sidebar."],
    ], [1.8, 4.7])

    heading(doc, "14.4 Desinstalación", 2)
    body(doc, "UNIRED es una aplicación web, por lo que los usuarios finales no requieren desinstalar nada. Para dejar de usar el sistema, simplemente dejen de acceder a la URL. Si desean eliminar sus datos, pueden usar la opción 'Eliminar Cuenta' en la sección de edición de perfil, lo cual eliminará permanentemente toda su información del sistema.")

    doc.add_paragraph()
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FIN DEL DOCUMENTO")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    add_hr(doc)


# ── Page Setup ─────────────────────────────────────────────────
def setup_page(doc):
    """Configure page margins."""
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def add_page_number(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("UNIRED — Documentación Técnica")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ── Main ───────────────────────────────────────────────────────
def build_document():
    doc = Document()

    # Document properties
    doc.core_properties.title = "Sistema Integral de Red Social Universitaria - UNIRED"
    doc.core_properties.author = "Gabriel Lauro Hernández, Jose Manuel Orozco Vazquez, David Gonzalez Vargas"
    doc.core_properties.subject = "Documentación Técnica y de Análisis - Ingeniería de Software II"

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    setup_page(doc)
    add_page_number(doc)

    # Build all sections
    build_cover(doc)
    build_feasibility(doc)
    build_interviews(doc)
    build_requirements_table(doc)
    build_functional_requirements(doc)
    build_validation(doc)
    build_analysis_unit(doc)
    build_use_cases(doc)
    build_design(doc)
    build_database(doc)
    build_sequence_diagrams(doc)
    build_class_diagrams(doc)
    build_testing(doc)
    build_deployment(doc)
    build_user_manual(doc)

    doc.save(OUTPUT)
    print(f"DOCX generated successfully: {OUTPUT}")
    print(f"File size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
